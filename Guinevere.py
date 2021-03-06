#!/usr/bin/python
#!/usr/bin/env python

"""Guinevere is a tool used to automate security assessment reporting"""

import MySQLdb, os, docx, argparse, math, netaddr, logging, readline, sys, json
from warnings import filterwarnings, resetwarnings
from docx.enum.text import WD_ALIGN_PARAGRAPH
from cvss import CVSS2

#Requires MySQL driver, python-mysqldb for Linux. Seems to be installed in Kali
#Requires python-docx library, apt-get update; apt-get install -y python-pip mysql-client;pip install python-docx
#Requires cvsss library; pip install cvss
### OSX Install Notes:
# sudo su -
# export CFLAGS=-Qunused-arguments
# export CPPFLAGS=-Qunused-arguments
# pip install mysql-python
# pip install python-docx
# pip install netaddr

#################################################
#           Guinevere Variables                 #
#################################################
__author__ = "Russel Van Tuyl"
__license__ = "GPL"
__version__ = "1.4.1"
__maintainer__ = "Russel Van Tuyl"
__email__ = "Russel.VanTuyl@gmail.com"
__status__ = "Development"
G_root = os.path.dirname(os.path.realpath(__file__))
readline.parse_and_bind('tab: complete')
readline.set_completer_delims('\t')
# logging.basicConfig(stream=sys.stdout, format='%(asctime)s\t%(levelname)s\t%(message)s',
                    # datefmt='%Y-%m-%d %I:%M:%S %p', level=logging.DEBUG)  # Log to STDOUT
logging.basicConfig(filename=os.path.join(G_root, 'Guinevere.log'), format='%(asctime)s\t%(levelname)s\t%(message)s',
                    datefmt='%Y-%m-%d %I:%M:%S %p', level=logging.DEBUG)  # Log to File
#################################################
#CHANGE TO MATCH YOUR DATABASE
g_ip = "127.0.0.1"          # Database IP address
g_p = 3306                  # Database Port
g_user = "gauntlet"             # Database Username
g_pass = "password"         # Database Password
#################################################
#                   COLORS                      #
#################################################
note = "\033[0;0;33m-\033[0m"
warn = "\033[0;0;31m!\033[0m"
info = "\033[0;0;36mi\033[0m"
question = "\033[0;0;37m?\033[0m"
debug = "\033[0;0;31m[DEBUG]\033[0m"

#Parse command line arguments
parser = argparse.ArgumentParser()
parser.add_argument('-H', '--db-host', type=str, default=g_ip, help="MySQL Database Host. Default set in script")
parser.add_argument('-U', '--db-user', type=str, default=g_user, help="MySQL Database Username. Default set in script")
parser.add_argument('-P', '--db-pass', type=str, default=g_pass, help="MySQL Database Password. Default set in script")
parser.add_argument('-p', '--db-port', type=str, default=g_p, help="MySQL Database Port. Default set in script")
parser.add_argument('-l', '--lines', type=int, default=10, help="Number of lines to display when selecting an "
                                                                "engagement. Default is 10")
parser.add_argument('-A', '--all-vulns', action='store_true', default=False, help="Include all vulnerability headings "
                                                                                  "when there are no associated report "
                                                                                  "narratives")
parser.add_argument('-V', '--all-verb', action='store_true', default=False, help="Include all vureto vulnerability "
                                                                                 "verbiage when there are no "
                                                                                 "associated report narratives")
parser.add_argument('--ports', action='store_false', default=True, help="Exclude port information vulnerability "
                                                                        "write-up portion of the report")
parser.add_argument('--cvss', action='store_false', default=True, help="Exclude CVSS scores from vulnerability titles")
parser.add_argument('-sC', action='store_false', default=True, help="Exclude Critical-Severity Vulnerabilities")
parser.add_argument('-sH', action='store_false', default=True, help="Exclude High-Severity Vulnerabilities")
parser.add_argument('-sM', action='store_false', default=True, help="Exclude Medium-Severity Vulnerabilities")
parser.add_argument('-sL', action='store_true', default=False, help="Include Low-Severity Vulnerabilities")
parser.add_argument('-sI', action='store_true', default=False, help="Include Informational-Severity Vulnerabilities")
parser.add_argument('-aD', '--assessment-date', action='store_true', default=False, help='Include the date when '
                                                                                         'selecting an assessment '
                                                                                         'to report on')
parser.add_argument('-T', '--tool-output', action='store_false', default=True, help="Exclude Tool Output When Printing "
                                                                                    "G-Checklist")
parser.add_argument('--debug', action='store_true', default=False, help="Enable debug output to console")
args = parser.parse_args()


def get_assessment(info_string_for_user):
    """Connects to the assessment database and prompts user to select an engagement"""

    logging.info('Entering the get_assessment function')
    db = MySQLdb.connect(host=args.db_host, user=args.db_user, passwd=args.db_pass, port=args.db_port)
    gauntlet=db.cursor()
    gauntlet.execute("""show databases""")
    all_available_databases = gauntlet.fetchall()
    gauntlet.close()
    discovered_gauntlet_databases = []
    os.system('clear')
    banner()

    # number of databases to print at a time
    database_print_limit_counter = args.lines

    # add all databases that start with 'gauntlet_' to a list
    for discovered_db in all_available_databases:
        if discovered_db[0].startswith('gauntlet'):
            discovered_gauntlet_databases.append(discovered_db[0].replace("gauntlet_", ''))

    #Print engagements to screen and have user choose one
    # Counter for number of engagements printed to the screen
    printed_engagement_counter = 0
    z = False
    # Counter for the number of engagements to display at a time
    eng_display_limit_counter = database_print_limit_counter
    while z != True:
        if (printed_engagement_counter <= eng_display_limit_counter) and (printed_engagement_counter < len(discovered_gauntlet_databases)):
            if args.assessment_date:
                SDate = db_query('select value  FROM engagement_details WHERE `key`="Start Date"', discovered_gauntlet_databases[printed_engagement_counter])
                if SDate is not None:
                    print "[" + str(printed_engagement_counter) + "]" + discovered_gauntlet_databases[printed_engagement_counter] + "\t" + SDate[0][0]
            else:
                print "[" + str(printed_engagement_counter) + "] " + discovered_gauntlet_databases[printed_engagement_counter]
            printed_engagement_counter += 1
        else:
            print "[99] More..."
            user_selection_string = raw_input("\n[" + question +"]Please select " + info_string_for_user + ": ")
            try:
                if ((user_selection_string == "99") or (user_selection_string == "")):
                    if printed_engagement_counter == len(discovered_gauntlet_databases):
                        printed_engagement_counter = 0
                        eng_display_limit_counter = database_print_limit_counter
                    else:
                        eng_display_limit_counter = eng_display_limit_counter + database_print_limit_counter
                    os.system('clear')
                    banner()
                elif user_selection_string == "Q" or user_selection_string == "q":
                    exit()
                elif discovered_gauntlet_databases[int(user_selection_string)]:
                    z = True
                else:
                    pass
            except:
                    os.system('clear')
                    banner()
                    print "["+warn+"]ERROR: " + user_selection_string + " is not a valid option. Try again"
                    printed_engagement_counter = 0
                    eng_display_limit_counter = database_print_limit_counter
    os.system('clear')
    logging.info(discovered_gauntlet_databases[int(user_selection_string)] + " assessment selected by user")

    return discovered_gauntlet_databases[int(user_selection_string)]


def get_crosstable(assessment):
    """Select the which assessment crosstable to use"""

    logging.info('Entering the get_crosstable function')
    chosen_crosstable = ""
    list_of_crosstables = ()
    tries = 0
    #Find all the crosstables for the assessment
    while not list_of_crosstables:
        try:
            if tries < 100:
                q = """SELECT DISTINCT table_id from cross_data_nva"""
                list_of_crosstables = db_query(q, assessment)
            else:
                print "Its broken try again"
                break
        except:
            tries = tries+1

    while not chosen_crosstable:

        # If there is more than 1 crosstable, have the user choose
        if len(list_of_crosstables) > 1:
            list_number = 0
            for crosstable in list_of_crosstables:
                print "[" + str(list_number) + "]",crosstable[0]
                list_number += 1
            try:
                user_choice = raw_input("\nWhich Crosstable would you like to use: ")
                chosen_crosstable = list_of_crosstables[int(user_choice)]
            except:
                os.system('clear')
                banner()
                print "["+warn+"]Error: please try again"

        # Otherwise choose for the user
        else:
            chosen_crosstable = list_of_crosstables[0]

    os.system('clear')
    banner()
    logging.info(str(chosen_crosstable[0]) + " crosstable selected")
    return chosen_crosstable[0]


def assessment_vulns(assessment, crosstable):
    """Builds a list of the assessment vulnerabilities"""

    vulns = []
    plugins = ""
    #Import data from gauntlet db for the selected crosstable
    hosts = db_query("select * from cross_data_nva WHERE table_id = '%s'" % (crosstable), assessment)

    # Vureto Vuln ID is in spot 2 of the tuple returned by i in the following loop
    Vureto_Vuln_ID_Location = 2

    for i in hosts:
        Vureto_Vuln_ID = i[Vureto_Vuln_ID_Location]
        if ((Vureto_Vuln_ID.startswith('Nessus') or Vureto_Vuln_ID.startswith('Netsparker') or Vureto_Vuln_ID.startswith('Acunetix')
                or Vureto_Vuln_ID.startswith('BurpSuite')) and Vureto_Vuln_ID not in plugins):
            plugins += Vureto_Vuln_ID
        else:
            vulns.append(Vureto_Vuln_ID)
    if plugins != "":
        pass
    return vulns


def assessment_report(vulnerability_id_info_mapping_list):
    """Builds a unique list of Report IDs for the selected assessment and crosstable"""

    logging.info('Entered assessment_report function')
    temp = []
    for vulnerability_id in vulnerability_id_info_mapping_list:
        if args.debug:
            logging.info('Vulnerability: ' + vulnerability_id + " - " + vulnerability_id_info_mapping_list[vulnerability_id]['vuln_title'])
        if vulnerability_id_info_mapping_list[vulnerability_id]['vuln_report_id'] is not None:
            temp.append(vulnerability_id_info_mapping_list[vulnerability_id]['vuln_report_id'])
        else:
            pass
    logging.info('Leaving assessment_report_function')
    return set(temp)


def get_vulns(vuln_IDs, assessment, crosstable):
    """Build dictionary containing the assessment vulnerabilities and their associated information"""

    logging.info('Entered get_vulns function...')
    vulns = {}
    plugins = ""
    tools = ['Nessus', 'Netsparker', 'Acunetix', 'BurpSuite', 'Nmap', 'Nikto', 'dirb']  # names of tools to ignore
    db = MySQLdb.connect(host=args.db_host, user=args.db_user, passwd=args.db_pass, port=args.db_port, db='GauntletData')
    filterwarnings('ignore', category = MySQLdb.Warning) # Disable MySQL Warnings

    # Variable for progress bar
    countvulnVar = 0

    # vuln_IDs is an array of vuln id's created in assessment_vulns function
    for vuln_id in vuln_IDs:

        # Create a progress bar
        countvulnVar = countvulnVar + 1.0
        progress = str(countvulnVar / len(vuln_IDs) * 100)
        print "\r\t[" + note + "]Querying Database For Vulnerability Information:", progress[:5] + "%",

        #TODO to remove "Nessus 1111" entries

        if vuln_id.split()[0] in tools:
            if vuln_id not in plugins and vuln_id is "":
                plugins += "\t["+warn+"]" + vuln_id + " plugin needs to be added to your Gauntlet database"
            elif vuln_id not in plugins:
                plugins += "\n\t["+warn+"]" + vuln_id + " plugin needs to be added to your Gauntlet database"

        # Get vulnerability verbage
        else:
            gauntlet = db.cursor()
            gauntlet.execute("""select title, description, solution, report_id from vulns WHERE gnaat_id=%s""", (vuln_id,))
            temp = gauntlet.fetchone()
            gauntlet.close()
            if temp[3] is not None:

                vulns[vuln_id] = {'vuln_id': vuln_id, 'vuln_title': temp[0], 'vuln_desc': temp[1], 'vuln_sol': temp[2], 'vuln_report_id': int(temp[3])}
            else:
                vulns[vuln_id] = {'vuln_id': vuln_id, 'vuln_title': temp[0], 'vuln_desc': temp[1], 'vuln_sol': temp[2], 'vuln_report_id': temp[3]}

    db2 = MySQLdb.connect(host=args.db_host, user=args.db_user, passwd=args.db_pass, port=args.db_port, db='gauntlet_'+ assessment)
    #Add all hosts with the associated vulnerability to the rpt dictionary
    print ""
    countvulnVar = 0

    #Determine if cross_data_nva has a 'port' column
    portColumn = False
    gauntlet=db2.cursor()
    gauntlet.execute("""SHOW COLUMNS FROM cross_data_nva FROM gauntlet_%s;""" % assessment)
    columns = gauntlet.fetchall()
    gauntlet.close
    for c in columns:
        if c[0] == 'port':
            portColumn = True

    for vuln_id in vuln_IDs:

        # Create progress bar
        countvulnVar = countvulnVar + 1.0
        progress = str(countvulnVar / len(vuln_IDs) * 100)
        print "\r\t[" + note + "]Querying Database For Affected Hosts:", progress[:5] + "%",

        # Ignore id's in tool list
        if vuln_id.split()[0] in tools:
            pass
        else:
            gauntlet=db2.cursor()
            if args.ports and portColumn:
                gauntlet.execute(
                    """SELECT host, port, protocol FROM cross_data_nva WHERE cross_data_nva.table_id =%s
                    AND vuln_id=%s AND (s1='Y' or s2='Y' or s3='Y' or s4='Y' or s5='Y')""",
                    (crosstable,vuln_id))
            else:
                gauntlet.execute(
                    """SELECT host FROM cross_data_nva WHERE cross_data_nva.table_id =%s AND vuln_id=%s
                    AND (s1='Y' or s2='Y' or s3='Y' or s4='Y' or s5='Y')""",
                    (crosstable, vuln_id))
            temp2 = gauntlet.fetchall()
            gauntlet.close()
            #print vuln[j]['vuln_title'], temp2 #DEBUG
            vulns[vuln_id].update({'vuln_hosts': temp2})

    #Determine the rank of the vulnerability
    print ""
    countvulnVar = 0
    for vuln_id in vuln_IDs:

        # Create progress bar
        countvulnVar = countvulnVar + 1.0
        progress = str(countvulnVar / len(vuln_IDs) * 100)
        print "\r\t[" + note + "]Querying Database For Vulnerability Rank:", progress[:5] + "%",

        # Craft assessment string for SQL query
        if assessment == "GauntletData":
            assessment_name = assessment
        else:
            assessment_name = "gauntlet_" + assessment

        # Ignore ids in tool list
        if vuln_id.split()[0] in tools:
            pass
        # Select severity values from cross_tables
        else:
            temp4 = db_query("""SELECT s1, s2, s3, s4, s5 FROM cross_data_nva WHERE table_id ='%s' AND vuln_id = %s;""" % (crosstable, vuln_id), assessment)
            severities = []
            for value in temp4:
                if value[0] is 'Y' and ('Critical' not in severities):
                    severities.append('Critical')
                elif value[1] is 'Y' and ('High' not in severities):
                    severities.append('High')
                elif value[2] is 'Y' and ('Medium' not in severities):
                    severities.append('Medium')
                elif value[3] is 'Y' and ('Low' not in severities):
                    severities.append('Low')
                elif value[4] is 'Y' and ('Informational' not in severities):
                    severities.append('Informational')
                else:
                    if None in severities:
                        pass
                    else:
                        severities.append(None)

            # Update vuln_rating to the highest discovered severity
            if 'Critical' in severities:
                vulns[vuln_id].update({'vuln_rating': 'Critical'})
            elif 'High' in severities:
                vulns[vuln_id].update({'vuln_rating': 'High'})
            elif 'Medium' in severities:
                vulns[vuln_id].update({'vuln_rating': 'Medium'})
            elif 'Low' in severities:
                vulns[vuln_id].update({'vuln_rating': 'Low'})
            elif 'Informational' in severities:
                vulns[vuln_id].update({'vuln_rating': 'Informational'})
            else:
                vulns[vuln_id].update({'vuln_rating': None})

    if plugins != "":
        print "%s" %(plugins,),
    print ""
    resetwarnings() #Re-enable MySQL Warnings
    return vulns


def get_report(report_record_IDs, vuln_ID_info_mapping):
    """Build a dictionary containing all of the reporting information"""
    logging.info("Entering get_report function")

    rpt = {}

    # change to GauntletData after dev/or vureto for dev
    db = MySQLdb.connect(host=args.db_host, user=args.db_user, passwd=args.db_pass, port=args.db_port, db='GauntletData')
    
    for report_record_ID in report_record_IDs:
        gauntlet = db.cursor()
        gauntlet.execute("""select title, identification, explanation, impact, recommendation from report
                         WHERE report_id=%s""" , (report_record_ID,))
        temp = gauntlet.fetchone()
        gauntlet.close()
        rpt[report_record_ID] = {'report_id': report_record_ID, 'report_title': temp[0], 'report_identification': temp[1],
                  'report_explanation': temp[2], 'report_impact': temp[3], 'report_recommendation': temp[4]}

    # Add all vulnerabilities with this report ID to the dictionary
    for vuln_ID in vuln_ID_info_mapping:
        if vuln_ID_info_mapping[vuln_ID]['vuln_report_id'] is not None:
            if 'vulns' in rpt[vuln_ID_info_mapping[vuln_ID]['vuln_report_id']]:
                rpt[vuln_ID_info_mapping[vuln_ID]['vuln_report_id']]['vulns'][vuln_ID_info_mapping[vuln_ID]['vuln_id']] = vuln_ID_info_mapping[vuln_ID]
            else:
                rpt[vuln_ID_info_mapping[vuln_ID]['vuln_report_id']]['vulns'] = {vuln_ID_info_mapping[vuln_ID]['vuln_id']: vuln_ID_info_mapping[vuln_ID]}
        else:
            pass

    # Determine the highest severity level and set it for the reporting record
    for record in rpt:
        r = []
        for vulnerability in rpt[record]['vulns']:
            r.append(rpt[record]['vulns'][vulnerability]['vuln_rating'])
        if 'Critical' in r:
            rpt[record]['report_rating'] = 'Critical'
            continue
        elif 'High' in r:
            rpt[record]['report_rating'] = 'High'
            continue
        elif 'Medium' in r:
            rpt[record]['report_rating'] = 'Medium'
            continue
        elif 'Low' in r:
            rpt[record]['report_rating'] = 'Low'
            continue
        elif 'Informational' in r:
            rpt[record]['report_rating'] = 'Informational'
            continue
        else:
            rpt[record]['report_rating'] = None
    logging.info('Leaving get_report Function')
    return rpt


def banner():
    """Guinevere's banner"""
    #Art retrieved from http://www.oocities.org/spunk1111/women.htm

    print """        ,,,_"""
    print """     .'     `'. ################################################"""
    print "    /     ____ \\#               Guinevere v"+__version__+"               #"
    print "   |    .`_  _\/#                                              #"
    print "   /    ) a  a| #   Automated Security Assessment Reporting    #"
    print "  /    (    > | ################################################"
    print """ (      ) ._  / """
    print " )    _/-.__.'`\\"
    print """(  .-'`-.   \__ )"""
    print """ `/      `-./  `.         """
    print "  |    \      \  \\"
    print "  |     \   \  \  \\"
    print "  |\     `. /  /   \\"
    print "_________________________________________________________________"


def int_to_string(i):
    """Converts an integer to its spelled out version; Used in reporting narratives"""
    
    s = { 0: "", #Dictionary of integers to spelled out words
        1: "one",  
        2: "two",
        3: "three",
        4: "four",
        5: "five",
        6: "six",
        7: "seven",
        8: "eight",
        9: "nine",
        10: "ten",
        11: "eleven",
        12: "twelve",
        13: "thirteen",
        14: "fourteen",
        15: "fifteen",
        16: "sixteen",
        17: "seventeen",
        18: "eighteen",
        19: "nineteen",
        20: "twenty",
        30: "thirty",
        40: "forty",
        50: "fifty",
        60: "sixty",
        70: "seventy",
        80: "eighty",
        90: "ninety"
        }

    #break i into ones, tens, hundreds, thousands and then build string
    #ONES
    if len(str(i)) is 1 and str(i) is "1":  #Spelling for host as opposed to hosts
        return s[i] +" (" + str(i) + ") host"
    elif len(str(i)) is 1 and str(i) is not "1": #Spelling for single digit hosts
        return s[i] +" (" + str(i) + ") hosts"
    #TENS
    elif len(str(i)) is 2 and str(i).startswith("1") :  #To grab spelling for 11 through 19
        return s[i]+" (" + str(i) + ") hosts"
    elif len(str(i)) is 2 and not str(i).startswith("1") and not str(i).endswith("0"): #To grab spelling for 20 through 99 where the number doesn't end in 0
        return s[(i/10)*10]+"-"+s[i -((i/10)*10)] +" (" + str(i) + ") hosts"
    elif len(str(i)) is 2 and not str(i).startswith("1") and str(i).endswith("0"): #To grab spelling for 20 through 99 where the number doesn't end in 0
        return s[(i/10)*10]+s[i -((i/10)*10)] +" (" + str(i) + ") hosts"
    #HUNDREDS
    elif len(str(i)) is 3 and "0" not in str(i): #to grab spelling for 100's where the number doesn't have a zero
        return s[(i/100)]+"-hundred "+s[((i -((i/100)*100))/10)*10]+"-"+s[i-(((i/100)*100)+((i -((i/100)*100))/10)*10)] +" (" + str(i) + ") hosts"
    elif len(str(i)) is 3 and not str(i).endswith("0") and "0" in str(i): #to grab spelling for 100's where the 10s place is 0
        return s[(i/100)]+"-hundred"+s[((i -((i/100)*100))/10)*10]+" "+s[i-(((i/100)*100)+((i -((i/100)*100))/10)*10)] +" (" + str(i) + ") hosts"
    elif len(str(i)) is 3 and str(i).endswith("0"): #to grab spelling for 100's where the number ends in 0
        return s[(i/100)]+"-hundred"+s[((i -((i/100)*100))/10)*10]+s[i-(((i/100)*100)+((i -((i/100)*100))/10)*10)] +" (" + str(i) + ") hosts"
    else:
        return "ERROR, was not able to return the number of host(s)"


def ip_sort(hosts):
    """Put the provided list of tuples IP addresses into order"""

    ips = []
    hostnames = []
    for ip in hosts:
        if unicode(ip[0].split('.')[0]).isnumeric():
            ips.append(netaddr.IPAddress(ip[0]))  # isnumeric only works with Unicode; checking for IP
        # elseif: ip[0].isalpha():         # Checking for when a hostname is used instead of an IP
        else:
            hostnames.append(ip[0])

    ips = sorted(ips)
    sorted_hosts = []
    for i in ips:                   # Add IPs
        sorted_hosts.append(str(i))
    for i in hostnames:             # Add Hostnames
        sorted_hosts.append(str(i))
    return sorted_hosts


def ip_sort_list(hosts):
    """Put the provided list IP addresses into order"""
    # TODO normalize both ip_sort functions to only use one data type (list or tuple)
    ips = []
    hostnames = []
    for ip in hosts:
        if unicode(ip.split('.')[0]).isnumeric():
            ips.append(netaddr.IPAddress(ip))  # isnumeric only works with Unicode; checking for IP
        else:
            hostnames.append(ip)

    ips = sorted(ips)
    sorted_hosts = []
    for i in ips:                   # Add IPs
        sorted_hosts.append(str(i))
    for i in hostnames:             # Add Hostnames
        sorted_hosts.append(str(i))
    return sorted_hosts


def generate_hosts_table(file, ass):
    """Build a list of assessment interesting hosts; hosts with atleast one TCP or UDP port open."""

    hosts = {}
    file.add_page_break()
    logging.info('Entering the generate_hosts_table function')
    print "["+note+"]Generating Interesting Hosts Table"
    # Build dictionary of host IDs and IPs from gauntlet's 'hosts' table
    engagement = db_query("""SELECT value FROM gauntlet_%s.engagement_details WHERE engagement_details.key = 'Engagement Task 1'""" % (ass), ass)
    if engagement:
        if 'Internal' in engagement[0][0]:
            temp = db_query("""SELECT host_id, ip_address, machine_name from hosts""", ass)
        else:
            temp = db_query("""SELECT host_id, ip_address, fqdn from hosts""", ass)
    else:
        temp = db_query("""SELECT host_id, ip_address, machine_name from hosts""", ass)
    temp2 = db_query("""SELECT host_id, port, protocol from ports""", ass)
    for i in temp:
        tcp = []
        udp = []
        for j in temp2:
            if (j[0] == i[0]) and (j[1] != "0"):
                if j[2] == 'tcp':
                    tcp.append(j[1])
                elif j[2] == 'udp':
                    udp.append(j[1])
                else:
                    pass
            else:
                pass
        hosts[i[0]] = {'IP': i[1], 'Name': i[2], 'TCP': tcp, 'UDP': udp}
    x = 0  # Number of interesting hosts counter
    for host in hosts:
        if (len(hosts[host]['TCP']) > 0) or (len(hosts[host]['UDP']) > 0):
            x += 1
        else:
            pass
    logging.info("["+info+"]"+str(x) + " Interesting Hosts")
    print "\t["+info+"]"+str(x) + " Interesting Hosts"
    file.add_heading(str(x) + ' Interesting Host(s) List')
    table = file.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'IP Address'
    hdr_cells[1].text = 'Hostname'
    hdr_cells[2].text = 'Open TCP Port(s)'
    hdr_cells[3].text = 'Open UDP Port(s)'
    table.style = 'Medium Grid 1 Accent 1'

    # Build a list of sorted IPs
    sorted_hosts = []
    for host in hosts:
        sorted_hosts.append(hosts[host]['IP'])
    sorted_hosts = ip_sort_list(sorted_hosts)

    for ip in sorted_hosts:
        for k in hosts:
            if hosts[k]['IP'] == ip:
                if (len(hosts[k]['TCP']) > 0) or (len(hosts[k]['UDP']) > 0):
                    x += 1
                    row_cells = table.add_row().cells
                    row_cells[0].text = hosts[k]['IP']
                    if len(hosts[k]['Name']) > 0:
                        row_cells[1].text = hosts[k]['Name']
                    else:
                        row_cells[1].text = "---"
                    if len(hosts[k]['TCP']) > 0:
                        row_cells[2].text = str(hosts[k]['TCP']).lstrip('[').rstrip(']').replace("'", "")
                    else:
                        row_cells[2].text = "---"
                    if len(hosts[k]['UDP']) > 0:
                        row_cells[3].text = str(hosts[k]['UDP']).lstrip('[').rstrip(']').replace("'", "")
                    else:
                        row_cells[3].text = "---"
                else:
                    pass

    return file


def generate_vuln_list(report_docx, assessment, vuln_dictionary_data):
    """Build the bullet list of vulnerabilities used in the executive summary"""

    logging.info("Entering the generate_vuln_list function")
    engagement = db_query("""SELECT value FROM gauntlet_%s.engagement_details WHERE engagement_details.key = 'Engagement Task 1'""" % (assessment), assessment)

    if engagement:
        report_docx.add_heading(str(engagement[0][0]) + ' NVA/PT')
    else:
        report_docx.add_heading('NVA/PT')

    def writeBullet(s, h):
        n = s.find('[n]') + 3  # Find '[n]' and add three to account for the length of '[n]'
        s = s.replace(s[0:n], int_to_string(h))
        s = s.rstrip('\n')
        s = s.rstrip('\t')
        s = s.rstrip()
        s = s.rstrip('.')
        s = s[0:1].upper() + s[1:]
        report_docx.add_paragraph(s, style='List Bullet')

    logging.info("Writing bullets based on criticality in the generate_vuln_list function")
    for i in vuln_dictionary_data:
        if args.debug:
            print "[" + info + "]%s" % vuln_dictionary_data[i]

        # Check to see if is a multi vuln report item
        if len(vuln_dictionary_data[i]['vulns']) > 1:
            h = 0
            for j in vuln_dictionary_data[i]['vulns']:
                h += len(vuln_dictionary_data[i]['vulns'][j]['vuln_hosts'])
            if args.sC and vuln_dictionary_data[i]['report_rating'] == 'Critical':
                writeBullet(vuln_dictionary_data[i]['report_identification'], h)
            elif args.sH and vuln_dictionary_data[i]['report_rating'] == 'High':
                writeBullet(vuln_dictionary_data[i]['report_identification'], h)
            elif args.sM and vuln_dictionary_data[i]['report_rating'] == 'Medium':
                writeBullet(vuln_dictionary_data[i]['report_identification'], h)
            elif args.sL and vuln_dictionary_data[i]['report_rating'] == 'Low':
                writeBullet(vuln_dictionary_data[i]['report_identification'], h)
            elif args.sI and vuln_dictionary_data[i]['report_rating'] == 'Informational':
                writeBullet(vuln_dictionary_data[i]['report_identification'], h)
            elif vuln_dictionary_data[i]['report_rating'] is None:
                print "\t[" + note + "]" + vuln_dictionary_data[i]['report_title'] + " has no affected hosts"
            else:
                pass

        else:
            for j in vuln_dictionary_data[i]['vulns']:
                if args.sC and vuln_dictionary_data[i]['report_rating'] == 'Critical':
                    writeBullet(vuln_dictionary_data[i]['report_identification'], len(vuln_dictionary_data[i]['vulns'][j]['vuln_hosts']))
                elif args.sH and vuln_dictionary_data[i]['report_rating'] == 'High':
                    writeBullet(vuln_dictionary_data[i]['report_identification'], len(vuln_dictionary_data[i]['vulns'][j]['vuln_hosts']))
                elif args.sM and vuln_dictionary_data[i]['report_rating'] == 'Medium':
                    writeBullet(vuln_dictionary_data[i]['report_identification'], len(vuln_dictionary_data[i]['vulns'][j]['vuln_hosts']))
                elif args.sL and vuln_dictionary_data[i]['report_rating'] == 'Low':
                    writeBullet(vuln_dictionary_data[i]['report_identification'], len(vuln_dictionary_data[i]['vulns'][j]['vuln_hosts']))
                elif args.sI and vuln_dictionary_data[i]['report_rating'] == 'Informational':
                    writeBullet(vuln_dictionary_data[i]['report_identification'], len(vuln_dictionary_data[i]['vulns'][j]['vuln_hosts']))
                elif vuln_dictionary_data[i]['report_rating'] is None:
                    print "\t[" + note + "]" + vuln_dictionary_data[i]['report_title'] + " has no affected hosts"
                else:
                    pass

    return report_docx


def db_query(q, assessment):
    """General use function used for querying the assessment database"""

    if assessment == "GauntletData":
        assessment2 = assessment
    else:
        assessment2 = "gauntlet_" + assessment

    db = MySQLdb.connect(host=args.db_host, user=args.db_user, passwd=args.db_pass, port=args.db_port, db=assessment2)
    try:
        gauntlet=db.cursor()
        gauntlet.execute(q)
        recordset = gauntlet.fetchall()
        gauntlet.close()
        return recordset
    except:
        print "There was an error performing the following query: "
        print q


def save_report(file, ass):
    """Save the generated assessment report"""
    out_dir = get_path()
    guinevere_file = os.path.join(out_dir, "Guinevere_"+ass+".docx")
    file.save(guinevere_file)
    print "["+warn+"]Report saved to: " + guinevere_file
    raw_input("["+question+"]Press enter to continue...")
    main_menu()


def retest():
    """Create a report for a retest of an assessment"""

    os.system('clear')
    banner()
    print "Retrieving available assessments..."

    # Collect data from the original assessment
    original_assessment = get_assessment("the original assessment")
    banner()
    original_crosstable = get_crosstable(original_assessment)
    print "["+note+"]Gathering original assessment vulnerability IDs..."
    original_vID = assessment_vulns(original_assessment, original_crosstable)
    print "["+note+"]Gathering original assessment vulnerability dataset..."
    original_vuln = get_vulns(original_vID, original_assessment, original_crosstable)

    # Collect data from the retest
    retest_assessment = get_assessment("the retest assessment")
    banner()
    retest_crosstable = get_crosstable(retest_assessment)
    print "["+note+"]Gathering retest vulnerability IDs..."
    retest_vID = assessment_vulns(retest_assessment, retest_crosstable)
    print "["+note+"]Gathering retest vulnerability dataset..."
    retest_vuln = get_vulns(retest_vID, retest_assessment, retest_crosstable)

    # Create the report stub
    retest_report = docx.Document()
    retest_report.add_heading(original_assessment+' Retest Results')

    retest = {} # Dictionary to hold retest data

    for i in original_vuln:
        if original_vuln[i]['vuln_rating'] is not None and original_vuln[i]['vuln_rating'] is not "Informational":
            retest[i] = {'vuln_id': i, 'vuln_title': original_vuln[i]['vuln_title'], 'vuln_rating': original_vuln[i]['vuln_rating'], 'total_orig': len(set(original_vuln[i]['vuln_hosts']))}
            if i in retest_vuln:
                o = set(original_vuln[i]['vuln_hosts']) #Original
                r = set(retest_vuln[i]['vuln_hosts'])   #Retest
                l = o - r                               #Leftover, fixed hosts
                b = []  # List of hosts from the original retest that are found in the retest

                for x in o: # For each host in the original assessment, check to see if it is in the retest assessment
                    if x in r:
                        b.append(x)
                if len(b) == 0:
                    print "\t["+note+"]" + original_vuln[i]['vuln_title'] + " - Remediated"
                    retest[i].update({'status': 'Remediated'})
                elif len(b) == len(o):
                    print "\t["+warn+"]" + original_vuln[i]['vuln_title'] + " - Not Remediated"
                    retest[i].update({'status': 'Not Remediated'})
                    retest[i].update({'v_hosts': o}) #Hosts Still Vulnerable, contributed by Zach
                else:
                    print "\t["+info+"]" + original_vuln[i]['vuln_title'] + \
                          " - Partially Remediated (Still vulnerable: " + str(len(b)) + ")"
                    retest[i].update({'status': 'Partially Remediated'})
                    retest[i].update({'v_hosts': b})#Hosts still vulnerable
                    retest[i].update({'f_hosts': l}) #Fixed hosts
            else:
                print "\t["+note+"]" + original_vuln[i]['vuln_title'] + " - Remediated"
                retest[i].update({'status': 'Remediated'})

    # Build Status Table
    retest_report.add_heading('Vulnerability Status')
    status_table = retest_report.add_table(rows=1, cols=3)
    status_table.style = 'Medium Grid 1 Accent 1'
    hdr_cells = status_table.rows[0].cells
    hdr_cells[0].text = 'Severity'
    hdr_cells[1].text = 'Vulnerability'
    hdr_cells[2].text = 'Status'

    # Add Critical first
    for i in retest:
        if retest[i]['vuln_rating'] is 'Critical':
            row_cells = status_table.add_row().cells
            row_cells[0].text = retest[i]['vuln_rating']
            row_cells[1].text = retest[i]['vuln_title']
            row_cells[2].text = retest[i]['status']

    # Add High second
    for i in retest:
        if retest[i]['vuln_rating'] is 'High':
            row_cells = status_table.add_row().cells
            row_cells[0].text = retest[i]['vuln_rating']
            row_cells[1].text = retest[i]['vuln_title']
            row_cells[2].text = retest[i]['status']

    # Add Medium third
    for i in retest:
        if retest[i]['vuln_rating'] is 'Medium':
            row_cells = status_table.add_row().cells
            row_cells[0].text = retest[i]['vuln_rating']
            row_cells[1].text = retest[i]['vuln_title']
            row_cells[2].text = retest[i]['status']

    # Add Low last
    for i in retest:
        if retest[i]['vuln_rating'] is 'Low':
            row_cells = status_table.add_row().cells
            row_cells[0].text = retest[i]['vuln_rating']
            row_cells[1].text = retest[i]['vuln_title']
            row_cells[2].text = retest[i]['status']

    # Build Still Vulnerable Hosts Table
    retest_report.add_heading('Hosts Still Vulnerable')
    vulnerable_table = retest_report.add_table(rows=1, cols=3)
    vulnerable_table.style = 'Medium Grid 1 Accent 1'
    hdr_cells = vulnerable_table.rows[0].cells
    hdr_cells[0].text = 'Severity'
    hdr_cells[1].text = 'Vulnerability'
    hdr_cells[2].text = 'Hosts'

    #Criticals
    for i in retest:
        # "and retest[i]['vuln_rating'] is not 'Informational' and len(retest[i]['v_hosts']) > 0" Contriubted by Zach
        if 'v_hosts' in retest[i] and retest[i]['vuln_rating'] is not 'Informational' and len(retest[i]['v_hosts']) > 0:
            if retest[i]['vuln_rating'] is 'Critical':
                row_cells = vulnerable_table.add_row().cells
                row_cells[0].text = retest[i]['vuln_rating']
                row_cells[1].text = retest[i]['vuln_title']
                hosts = []
                for h in retest[i]['v_hosts']:
                    hosts.append(h[0])
                row_cells[2].text = ((str(ip_sort_list(hosts)).replace("'", "")).lstrip("[")).rstrip("]")

    #Highs
    for i in retest:
        # "and retest[i]['vuln_rating'] is not 'Informational' and len(retest[i]['v_hosts']) > 0" Contriubted by Zach
        if 'v_hosts' in retest[i] and retest[i]['vuln_rating'] is not 'Informational' and len(retest[i]['v_hosts']) > 0:
            if retest[i]['vuln_rating'] is 'High':
                row_cells = vulnerable_table.add_row().cells
                row_cells[0].text = retest[i]['vuln_rating']
                row_cells[1].text = retest[i]['vuln_title']
                hosts = []
                for h in retest[i]['v_hosts']:
                    hosts.append(h[0])
                row_cells[2].text = ((str(ip_sort_list(hosts)).replace("'", "")).lstrip("[")).rstrip("]")

    #Mediums
    for i in retest:
        # "and retest[i]['vuln_rating'] is not 'Informational' and len(retest[i]['v_hosts']) > 0" Contriubted by Zach
        if 'v_hosts' in retest[i] and retest[i]['vuln_rating'] is not 'Informational' and len(retest[i]['v_hosts']) > 0:
            if retest[i]['vuln_rating'] is 'Medium':
                row_cells = vulnerable_table.add_row().cells
                row_cells[0].text = retest[i]['vuln_rating']
                row_cells[1].text = retest[i]['vuln_title']
                hosts = []
                for h in retest[i]['v_hosts']:
                    hosts.append(h[0])
                row_cells[2].text = ((str(ip_sort_list(hosts)).replace("'", "")).lstrip("[")).rstrip("]")

    #Lows
    for i in retest:
        # "and retest[i]['vuln_rating'] is not 'Informational' and len(retest[i]['v_hosts']) > 0" Contriubted by Zach
        if 'v_hosts' in retest[i] and retest[i]['vuln_rating'] is not 'Informational' and len(retest[i]['v_hosts']) > 0:
            if retest[i]['vuln_rating'] is 'Low':
                row_cells = vulnerable_table.add_row().cells
                row_cells[0].text = retest[i]['vuln_rating']
                row_cells[1].text = retest[i]['vuln_title']
                hosts = []
                for h in retest[i]['v_hosts']:
                    hosts.append(h[0])
                row_cells[2].text = ((str(ip_sort_list(hosts)).replace("'", "")).lstrip("[")).rstrip("]")


    # Build stats table
    o_total_c = 0   # Original Total Critical
    r_total_c = 0   # Retest Total Critical
    o_total_h = 0
    r_total_h = 0
    o_total_m = 0
    r_total_m = 0
    o_total_l = 0
    r_total_l = 0
    for i in retest:
        # Critical Vulnerabilities
        if retest[i]['vuln_rating'] is 'Critical':
            o_total_c += retest[i]['total_orig']
            if 'v_hosts' in retest[i]:
                r_total_c += len(retest[i]['v_hosts'])
        # High Vulnerabilities
        if retest[i]['vuln_rating'] is 'High':
            o_total_h += retest[i]['total_orig']
            if 'v_hosts' in retest[i]:
                r_total_h += len(retest[i]['v_hosts'])
        # Medium Vulnerabilities
        if retest[i]['vuln_rating'] is 'Medium':
            o_total_m += retest[i]['total_orig']
            if 'v_hosts' in retest[i]:
                r_total_m += len(retest[i]['v_hosts'])
        # Low Vulnerabilities
        if retest[i]['vuln_rating'] is 'Low':
            o_total_l += retest[i]['total_orig']
            if 'v_hosts' in retest[i]:
                r_total_l += len(retest[i]['v_hosts'])

    s = "The original security assessment identified (" + str(o_total_c) + ") critical-severity, (" + str(o_total_h) + ") high-severity, (" + str(o_total_m) + ") medium-severity, and (" + str(o_total_l) + ") low-severity vulnerabilities."

    # Setup Table
    retest_report.add_heading('Retest Statistics')
    retest_report.add_paragraph(s)
    stats_table = retest_report.add_table(rows=1, cols=5)
    stats_table.style = 'Medium Grid 1 Accent 1'
    hdr_cells = stats_table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Critical'
    hdr_cells[2].text = 'High'
    hdr_cells[3].text = 'Medium'
    hdr_cells[4].text = 'Low'
    # Original Assessment Numbers
    row_cells = stats_table.add_row().cells
    row_cells[0].text = 'Original'
    row_cells[1].text = str(o_total_c)
    row_cells[2].text = str(o_total_h)
    row_cells[3].text = str(o_total_m)
    row_cells[4].text = str(o_total_l)

    # Retest Assessment Numbers
    row_cells = stats_table.add_row().cells
    row_cells[0].text = 'Retest'
    row_cells[1].text = str(r_total_c)
    row_cells[2].text = str(r_total_h)
    row_cells[3].text = str(r_total_m)
    row_cells[4].text = str(r_total_l)

    save_report(retest_report, retest_assessment)


def main_menu():
    """Display the main menu"""

    logging.info('Entered into main_menu function')
    i = None
    valid_options = {1: generate_assessment_report,
                     2: sql_dump,
                     3: retest,
                     4: patch_gauntlet,
                     5: pentest_checklist,
                     6: generate_assessment_json,
                     7: exit,
    }
    os.system('clear')
    banner()
    try:
        while i is None:
            print "\t\t\t\033[0;0;37mGUINEVERE MAIN MENU\033[0m\n"
            print "[1]Generate Assessment Report"
            print "[2]Export Assessment"
            print "[3]Generate Retest Report"
            print "[4]Patch Gauntled Database"
            print "[5]Generate Pentest Checklist"
            print "[6]Generate Assessment JSON File"
            print "[7]Exit"
            i = raw_input("\nWhat would you like to do: ")
            if int(i) in valid_options:
                valid_options[int(i)]()
            else:
                os.system('clear')
                banner()
                print "["+warn+"]" + str(i) + " is not a valid option, please try again: "
                i = None
    except ValueError:
        main_menu()


def sql_dump():
    """Use mysqldump to export an assessment to a .sql file"""

    import subprocess, time, platform

    #Find the script location and set the mysqldump executable
    #Will implement functionality in the future
    script_path = os.path.dirname(os.path.realpath(__file__))
    bin_path = os.path.join(script_path, 'bin')
    #Determine the operating system
    operating_system = platform.platform()
    if 'Windows' in operating_system:
        ext = '.exe'
    elif 'Linux' in operating_system:
        ext = '.bin'
    else:
        ext = '.nothing'
    mysqldump = "mysqldump"
    if os.path.isdir(bin_path) and os.path.exists(bin_path+'mysqldump.'+ext):
        mysqldump = bin_path+'mysqldump.'+ext

    assessment = get_assessment("the assessment to backup")
    os.system('clear')
    banner()
    output_path = get_path()
    date_time = time.strftime('%m%d%Y-%H%M%S')
    try:
        sql_file = open(os.path.join(output_path, assessment+"_"+date_time+".sql"), "w")
        subprocess.call([mysqldump, "--host="+args.db_host, "-u", args.db_user, "-p"+args.db_pass, "gauntlet_"+assessment], stdout=sql_file)
        #os.system(mysqldump+" --host="+args.db_host+" -u "+args.db_user+" -p"+args.db_pass+" gauntlet_"+assessment+" > "+sql_file)
        print "["+warn+"]SQL file saved to: " + os.path.join(output_path, assessment+"_"+date_time+".sql")
    except OSError:
        print "["+warn+"]mysqldump is likely not in your path, please add it and try again"
        raise
    except:
        raise #Just use for debug
    raw_input("["+question+"]Press enter to continue...")
    main_menu()


def get_path():
    """Prompt the user to enter a directory path"""

    output_path = None
    while output_path is None:
        print "["+question+"]Please enter the directory where you would like the file saved?"
        output_path = raw_input()
        if os.path.isdir(os.path.expanduser(output_path)):
            pass
        else:
            os.system('clear')
            banner()
            print "["+warn+"]" + str(output_path) + " is not valid, please try again: "
            output_path = None
    return os.path.expanduser(output_path)


def write_single_vul(rpt, report):
    """Write the single vulnerability paragraph"""

    cvss = None
    if len(rpt['vulns']) is 1:
        for r in rpt['vulns']:
            cvss = get_cvss(r)

    if args.cvss:
        report.add_heading("%s (CVSS: %0.1f - %s)" % (rpt['report_title'], cvss.base_score, rpt['report_rating']),
                           level=4)
    else:
        report.add_heading("%s (%s)" % (rpt['report_title'], rpt['report_rating']), level=4)

    for i in rpt['vulns']:
        vulnerableHosts = []
        for j in set(rpt['vulns'][i]['vuln_hosts']):  # Build Single Dimensional List
            vulnerableHosts.append(j)
        sortedHosts = ip_sort(vulnerableHosts)  # Create a unique & sorted list of hosts (avoids duplicate hosts)
        hosts = []
        for h in sortedHosts:
            for h2 in vulnerableHosts:
                if h2[0] == h:
                    if args.ports:
                        if len(h2) == 3:
                            if h2[1] != '0' and h2[2] != 'icmp':
                                host_info = h2[0] + ":" + h2[1] + "/" + h2[2]
                                if host_info not in hosts:
                                    hosts.append(host_info)
                            else:
                                if h2[0] not in hosts:
                                    hosts.append(h2[0])
                        else:
                            if h2[0] not in hosts:
                                hosts.append(h2[0])
                    else:
                        if h2[0] not in hosts:
                            hosts.append(h2[0])

        if len(hosts) == 1:         # If there is just one host
            p = rpt['report_identification'].replace("[n]", int_to_string(len(hosts))+ ' ('+hosts[0]+')')
        elif len(hosts) == 2:       # If there are two hosts
            p = rpt['report_identification'].replace("[n]", int_to_string(len(hosts))+ ' ('+hosts[0]+' and '+hosts[1]+')')
        elif len(hosts) >= 2 and len(hosts) <=5:  # If there are more than two but lest than five hosts
            host_list = ""
            for h in hosts:
                if h is hosts[len(hosts)-1]:  # Check to see if this is the last item in the list
                    host_list += "and " + h + ") "
                else:
                    host_list += h + ", "
            p = rpt['report_identification'].replace("[n]", int_to_string(len(hosts)) + ' ('+host_list)
        elif len(hosts) >= 6:  # If there are six or more hosts
            p = rpt['report_identification'].replace("[n]", int_to_string(len(hosts)) + '(refer to TABLE X)')

        if p.endswith(" ") or rpt['report_explanation'].startswith(" "):
            p += rpt['report_explanation']
        else:
            p += (" " + rpt['report_explanation'])
        if p.endswith(" ") or rpt['report_impact'].startswith(" "):
            p += rpt['report_impact']
        else:
            p += (" " + rpt['report_impact'])
        p = report.add_paragraph(p, style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p = report.add_paragraph(rpt['report_recommendation'], style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if args.debug:
            print debug + "Vulnerable Hosts: %s" % hosts
            raw_input(debug + "Press any key to continue")
        if len(hosts) >= 6:     # Draw the table
            c = 4  # number of desired columns
            r = int(math.ceil((len(hosts) / float(4))))  # Determine number of rows for table using a max of 4 columns
            hosts_table = report.add_table(rows=r, cols=c)
            hosts_table.style = 'Medium Grid 1 Accent 1'
            z = 0   # number of hosts
            x = 0   # row indices
            y = 0   # column indices
            while z < len(hosts):
                if (y / float(c)) == 1:  # Determine if we need to start putting data on a new row
                    y = 0   # reset column indices since max number of columns reached
                    x += 1
                hosts_table.cell(x, y).text = hosts[z]
                z += 1
                y += 1  # Add one to up the column data is put in
            if len(hosts)/float(c) != 1.000:  # Add "---" for empty spots in table
                d = c * (x+1)
                while d > len(hosts):
                    hosts_table.cell(x, y).text = "---"
                    d -= 1
                    y += 1

    return report


def write_multi_vul(rpt, report):
    """Write report data for grouped or multi vulnerabilities"""

    total_hosts = 0
    for i in rpt['vulns']:
        if args.sC and rpt['vulns'][i]['vuln_rating'] == 'Critical':
            total_hosts += len(rpt['vulns'][i]['vuln_hosts'])
        elif args.sH and rpt['vulns'][i]['vuln_rating'] == 'High':
            total_hosts += len(rpt['vulns'][i]['vuln_hosts'])
        elif args.sM and rpt['vulns'][i]['vuln_rating'] == 'Medium':
            total_hosts += len(rpt['vulns'][i]['vuln_hosts'])
        elif args.sL and rpt['vulns'][i]['vuln_rating'] == 'Low':
            total_hosts += len(rpt['vulns'][i]['vuln_hosts'])
        elif args.sI and rpt['vulns'][i]['vuln_rating'] == 'Informational':
            total_hosts += len(rpt['vulns'][i]['vuln_hosts'])

    report.add_heading(rpt['report_title'] + " (" + rpt['report_rating']+")", level=4)
    p = rpt['report_identification'].replace("[n]", int_to_string(total_hosts))
    p = p.rstrip('\n')
    p = p.rstrip('\t')
    p = p.rstrip()

    if p.endswith(" ") or rpt['report_explanation'].startswith(" "):
        p += rpt['report_explanation']
    else:
        p += (" " + rpt['report_explanation'])
    if p.endswith(" ") or rpt['report_impact'].startswith(" "):
        p += rpt['report_impact']
    else:
        p += (" " + rpt['report_impact'])

    p = report.add_paragraph(p, style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = report.add_paragraph(rpt['report_recommendation'], style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    table = None
    hdr_cells = None

    if args.cvss:
        table = report.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Severity'
        hdr_cells[1].text = 'CVSS'
        hdr_cells[2].text = 'Vulnerability'
        hdr_cells[3].text = 'Affected Host(s)'
    else:
        table = report.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Severity'
        hdr_cells[1].text = 'Vulnerability'
        hdr_cells[2].text = 'Affected Host(s)'

    table.style = 'Medium Grid 1 Accent 1'

    def writeRow(r, t, h, c=None):
        row_cells = table.add_row().cells
        if c is not None:
            row_cells[0].text = r
            row_cells[1].text = c
            row_cells[2].text = t
            row_cells[3].text = h
        else:
            row_cells[0].text = r
            row_cells[1].text = t
            row_cells[2].text = h

    for i in rpt['vulns']:
        if args.sC and rpt['vulns'][i]['vuln_rating'] == 'Critical':
            cvss = get_cvss(i)
            if args.cvss:
                writeRow(rpt['vulns'][i]['vuln_rating'], rpt['vulns'][i]['vuln_title'],
                         str(len(rpt['vulns'][i]['vuln_hosts'])), "%0.1f" % cvss.base_score)
            else:
                writeRow(rpt['vulns'][i]['vuln_rating'], rpt['vulns'][i]['vuln_title'],
                         str(len(rpt['vulns'][i]['vuln_hosts'])))
    for i in rpt['vulns']:
        if args.sH and rpt['vulns'][i]['vuln_rating'] == 'High':
            cvss = get_cvss(i)
            if args.cvss:
                writeRow(rpt['vulns'][i]['vuln_rating'], rpt['vulns'][i]['vuln_title'],
                         str(len(rpt['vulns'][i]['vuln_hosts'])), "%0.1f" % cvss.base_score)
            else:
                writeRow(rpt['vulns'][i]['vuln_rating'], rpt['vulns'][i]['vuln_title'],
                         str(len(rpt['vulns'][i]['vuln_hosts'])))
    for i in rpt['vulns']:
        if args.sM and rpt['vulns'][i]['vuln_rating'] == 'Medium':
            cvss = get_cvss(i)
            if args.cvss:
                writeRow(rpt['vulns'][i]['vuln_rating'], rpt['vulns'][i]['vuln_title'],
                         str(len(rpt['vulns'][i]['vuln_hosts'])), "%0.1f" % cvss.base_score)
            else:
                writeRow(rpt['vulns'][i]['vuln_rating'], rpt['vulns'][i]['vuln_title'],
                         str(len(rpt['vulns'][i]['vuln_hosts'])))
    for i in rpt['vulns']:
        if args.sL and rpt['vulns'][i]['vuln_rating'] == 'Low':
            cvss = get_cvss(i)
            if args.cvss:
                writeRow(rpt['vulns'][i]['vuln_rating'], rpt['vulns'][i]['vuln_title'],
                         str(len(rpt['vulns'][i]['vuln_hosts'])), "%0.1f" % cvss.base_score)
            else:
                writeRow(rpt['vulns'][i]['vuln_rating'], rpt['vulns'][i]['vuln_title'],
                         str(len(rpt['vulns'][i]['vuln_hosts'])))
    for i in rpt['vulns']:
        if args.sI and rpt['vulns'][i]['vuln_rating'] == 'Informational':
            cvss = get_cvss(i)
            if args.cvss:
                writeRow(rpt['vulns'][i]['vuln_rating'], rpt['vulns'][i]['vuln_title'],
                         str(len(rpt['vulns'][i]['vuln_hosts'])), "%0.1f" % cvss.base_score)
            else:
                writeRow(rpt['vulns'][i]['vuln_rating'], rpt['vulns'][i]['vuln_title'],
                         str(len(rpt['vulns'][i]['vuln_hosts'])))

    return report


def write_all_vuln(vuln, the_Report):

    print "["+note+"]Writing list of all vulnerabilities to the report: "
    the_Report.add_page_break()
    the_Report.add_heading("List of Assessment Vulnerabilities", 1)
    for i in vuln:
        if args.sC and vuln[i]['vuln_rating'] is 'Critical':
            print "\t["+info+"]"+ vuln[i]['vuln_title'], "(" + vuln[i]['vuln_rating'] + ")"
            the_Report.add_heading(vuln[i]['vuln_title'] + " (" + vuln[i]['vuln_rating'] + ")", 3)
            if args.all_verb:
                the_Report.add_paragraph(vuln[i]['vuln_desc'], style='Body Text')
                the_Report.add_paragraph(vuln[i]['vuln_sol'], style='Body Text')
    for i in vuln:
        if args.sH and vuln[i]['vuln_rating'] is 'High':
            print "\t["+info+"]"+ vuln[i]['vuln_title'], "(" + vuln[i]['vuln_rating'] + ")"
            the_Report.add_heading(vuln[i]['vuln_title'] + " (" + vuln[i]['vuln_rating'] + ")", 3)
            if args.all_verb:
                the_Report.add_paragraph(vuln[i]['vuln_desc'], style='Body Text')
                the_Report.add_paragraph(vuln[i]['vuln_sol'], style='Body Text')
    for i in vuln:
        if args.sM and vuln[i]['vuln_rating'] is 'Medium':
            print "\t["+info+"]"+ vuln[i]['vuln_title'], "(" + vuln[i]['vuln_rating'] + ")"
            the_Report.add_heading(vuln[i]['vuln_title'] + " (" + vuln[i]['vuln_rating'] + ")", 3)
            if args.all_verb:
                the_Report.add_paragraph(vuln[i]['vuln_desc'], style='Body Text')
                the_Report.add_paragraph(vuln[i]['vuln_sol'], style='Body Text')
    for i in vuln:
        if args.sL and vuln[i]['vuln_rating'] is 'Low':
            print "\t["+info+"]"+ vuln[i]['vuln_title'], "(" + vuln[i]['vuln_rating'] + ")"
            the_Report.add_heading(vuln[i]['vuln_title'] + " (" + vuln[i]['vuln_rating'] + ")", 3)
            if args.all_verb:
                the_Report.add_paragraph(vuln[i]['vuln_desc'], style='Body Text')
                the_Report.add_paragraph(vuln[i]['vuln_sol'], style='Body Text')
    for i in vuln:
        if args.sI and vuln[i]['vuln_rating'] is 'Informational':
            print "\t["+info+"]"+ vuln[i]['vuln_title'], "(" + vuln[i]['vuln_rating'] + ")"
            the_Report.add_heading(vuln[i]['vuln_title'] + " (" + vuln[i]['vuln_rating'] + ")", 3)
            if args.all_verb:
                the_Report.add_paragraph(vuln[i]['vuln_desc'], style='Body Text')
                the_Report.add_paragraph(vuln[i]['vuln_sol'], style='Body Text')
        #if vuln[i]['vuln_report_id'] is None and (((vuln[i]['vuln_rating'] is "Critical") and args.sC) or ((vuln[i]['vuln_rating'] is "High") and args.sH) or ((vuln[i]['vuln_rating'] is "Medium") and args.sM) or ((vuln[i]['vuln_rating'] is "Low") and args.sL) or ((vuln[i]['vuln_rating'] is "Informational") and args.sI)):
            #print "\t["+info+"]"+ vuln[i]['vuln_title'], "(" + vuln[i]['vuln_rating'] + ")"
            #the_Report.add_heading(vuln[i]['vuln_title'] + " (" + vuln[i]['vuln_rating'] + ")", 3)
            #if args.all_verb:
                #the_Report.add_paragraph(vuln[i]['vuln_desc'], style='BodyText')
                #the_Report.add_paragraph(vuln[i]['vuln_sol'], style='BodyText')

    return the_Report


def generate_assessment_report():
    """The main function for automatically generating an assessment report"""

    logging.info('Entering the generate_assessment_report function')
    os.system('clear')
    banner()
    print "Retrieving available assessments..."

    # Database containing the options of crosstables
    assessment = get_assessment("the assessment to create a report for")
    banner()

    # Choose crosstable from the options listed from get_assessment()
    crosstable = get_crosstable(assessment)

    vulnerability_ID_list = assessment_vulns(assessment, crosstable)
    os.system('clear')
    banner()
    print "["+note+"]Building list of found vulnerabilities for " + assessment + " Crosstable " + crosstable + "..."
    vuln_ID_to_info_mapping = get_vulns(vulnerability_ID_list, assessment, crosstable)

    ## FOLLOW vuln_id_to_info_mapping - Contains the severity rating, needs to be reset after each crosstable

    print "["+note+"]Generating report for the following vulnerabilities:"

    # Unique identifier for reporting record, allows access to verbage
    report_record_ID_set = assessment_report(vuln_ID_to_info_mapping)

    # Returns huge dictionary of vulnerabilities and associated data
    assessment_db = get_report(report_record_ID_set, vuln_ID_to_info_mapping)

    # Make the word document
    the_Report = docx.Document()
    the_Report.add_heading(assessment, 1)
    the_Report = generate_vuln_list(the_Report, assessment, assessment_db)

    if ((len(assessment_db) is 0) and args.all_vulns is False):
        exit("["+warn+"]Nothing to report on, quitting...")

    ####################################
    # Write the report in severity order
    ####################################

    logging.info("Writing the critical vulnerability narratives to the report")

    for i in assessment_db:
        if assessment_db[i]['report_rating'] == 'Critical' and args.sC:

            # Grouped Vulnerability Write-up
            if len(assessment_db[i]['vulns']) > 1:
                print '\t['+info+']Multi finding: ', assessment_db[i]['report_title']
                logging.info('['+info+']Multi finding: ' + assessment_db[i]['report_title'])
                the_report = write_multi_vul(assessment_db[i], the_Report)

            # Single Vulnerability Write-up
            elif assessment_db[i]['report_rating'] is not None:
                logging.info("["+info+"]" + assessment_db[i]['report_title'] +
                             "(" + assessment_db[i]['report_rating'] + ")")
                print "\t["+info+"]" + assessment_db[i]['report_title'] + " \033[0;0;31m(" + \
                    assessment_db[i]['report_rating'] + ")\033[0m"
                the_Report = write_single_vul(assessment_db[i], the_Report)

    logging.info("Writing the high vulnerability narratives to the report")

    for i in assessment_db:
        if assessment_db[i]['report_rating'] == 'High' and args.sH:
            if len(assessment_db[i]['vulns']) > 1:
                logging.info('['+info+']Multi finding: ' + assessment_db[i]['report_title'])
                print '\t['+info+']Multi finding: ', assessment_db[i]['report_title']
                the_report = write_multi_vul(assessment_db[i], the_Report)
            elif assessment_db[i]['report_rating'] is not None:
                logging.info("["+info+"]" + assessment_db[i]['report_title'] +
                             "(" + assessment_db[i]['report_rating'] + ")")
                print "\t["+info+"]" + assessment_db[i]['report_title'] + " \033[0;0;35m(" + \
                      assessment_db[i]['report_rating'] + ")\033[0m"
                the_Report = write_single_vul(assessment_db[i], the_Report)

    logging.info("Writing the medium vulnerability narratives to the report")

    for i in assessment_db:
        if assessment_db[i]['report_rating'] == 'Medium' and args.sM:
            if len(assessment_db[i]['vulns']) > 1:
                logging.info('['+info+']Multi finding: ' + assessment_db[i]['report_title'])
                print '\t['+info+']Multi finding: ', assessment_db[i]['report_title']
                the_report = write_multi_vul(assessment_db[i], the_Report)
            elif assessment_db[i]['report_rating'] is not None:
                logging.info("["+info+"]" + assessment_db[i]['report_title'] +
                             "(" + assessment_db[i]['report_rating'] + ")")
                print "\t["+info+"]" + assessment_db[i]['report_title'] + " \033[0;0;33m(" + \
                      assessment_db[i]['report_rating'] + ")\033[0m"
                the_Report = write_single_vul(assessment_db[i], the_Report)

    logging.info("Writing the low vulnerability narratives to the report")

    for i in assessment_db:
        if assessment_db[i]['report_rating'] == 'Low' and args.sL:
            if len(assessment_db[i]['vulns']) > 1:
                logging.info('['+info+']Multi finding: ' + assessment_db[i]['report_title'])
                print '\t['+info+']Multi finding: ', assessment_db[i]['report_title']
                the_report = write_multi_vul(assessment_db[i], the_Report)
            elif assessment_db[i]['report_rating'] is not None:
                logging.info("["+info+"]" + assessment_db[i]['report_title'] +
                             "(" + assessment_db[i]['report_rating'] + ")")
                print "\t["+info+"]" + assessment_db[i]['report_title'] + " \033[0;0;34m(" + \
                      assessment_db[i]['report_rating'] + ")\033[0m"
                the_Report = write_single_vul(assessment_db[i], the_Report)

    logging.info("Writing the informational vulnerability narratives to the report")

    for i in assessment_db:
        if assessment_db[i]['report_rating'] == 'Informational' and args.sI:
            if len(assessment_db[i]['vulns']) > 1:
                logging.info('['+info+']Multi finding: ' + assessment_db[i]['report_title'])
                print '\t['+info+']Multi finding: ', assessment_db[i]['report_title']
                the_report = write_multi_vul(assessment_db[i], the_Report)
            elif assessment_db[i]['report_rating'] is not None:
                logging.info("["+info+"]" + assessment_db[i]['report_title'] +
                             "(" + assessment_db[i]['report_rating'] + ")")
                print "\t["+info+"]" + assessment_db[i]['report_title'] + " \033[0;0;37m(" + \
                      assessment_db[i]['report_rating'] + ")\033[0m"
                the_Report = write_single_vul(assessment_db[i], the_Report)

    if args.all_vulns:
        the_Report = write_all_vuln(vuln_ID_to_info_mapping, the_Report)
    the_Report = generate_hosts_table(the_Report, assessment)
    save_report(the_Report, assessment)


def generate_assessment_json():
    """Generate a JSON object of the aggregated assessment information"""

    logging.info('Entering the generate_assessment_json function')
    os.system('clear')
    banner()
    print "Retrieving available assessments..."
    assessment = get_assessment("the assessment to create a JSON object for")
    banner()
    crosstable = get_crosstable(assessment)
    vID = assessment_vulns(assessment, crosstable)
    os.system('clear')
    banner()
    print "["+note+"]Building JSON object for " + assessment + " Crosstable " + crosstable + "..."
    vuln = get_vulns(vID, assessment, crosstable)
    rID = assessment_report(vuln)
    assessment_db = get_report(rID, vuln)
    engagment_details = gather_assessment_details(assessment)
    json_dict = {'engagment_details': engagment_details, 'report': assessment_db}
    json_object = json.dumps(json_dict)
    out_dir = get_path()
    json_file = os.path.join(out_dir, "Guinevere_" + assessment + "_" + crosstable + ".json")
    with open(json_file, "w") as j:
        j.write(json_object)
    print "["+warn+"]Assessment JSON object saved to: " + json_file
    raw_input("["+question+"]Press enter to continue...")
    main_menu()


def gather_assessment_details(assessment):
    """Gather assessment details from Gauntlet and create a dictionary"""

    logging.info('Entering the gather_assessment_details function')
    engagement = db_query("""SELECT value FROM gauntlet_%s.engagement_details WHERE
                          engagement_details.key = 'Engagement Task 1'""" % (assessment), assessment)[0][0]
    start_date = db_query("""SELECT value FROM gauntlet_%s.engagement_details WHERE
                          engagement_details.key = 'Start Date'""" % (assessment), assessment)[0][0]
    end_date = db_query("""SELECT value FROM gauntlet_%s.engagement_details WHERE
                          engagement_details.key = 'End Date'""" % (assessment), assessment)[0][0]
    analyst = db_query("""SELECT value FROM gauntlet_%s.engagement_details WHERE
                          engagement_details.key = 'Analyst 1'""" % (assessment), assessment)[0][0]

    i = {'engagment_type': engagement, 'start_date': start_date, 'stop_date': end_date, 'analyst': analyst}
    return i


def patch_gauntlet():
    print "Nothing to test right now"
    db = MySQLdb.connect(host=args.db_host, user=args.db_user, passwd=args.db_pass, port=args.db_port, db='GauntletData')

    create_table = """
        CREATE TABLE report (
            report_id integer NOT NULL AUTO_INCREMENT,
            title character varying(255) NOT NULL DEFAULT '',
            identification blob,
            explanation blob,
            impact blob,
            recommendation blob,
            status ENUM('NEW','MODIFIED','ACCEPTED','MARKED','DELETED') NOT NULL,
            owner character varying(255) NOT NULL DEFAULT '',
            PRIMARY KEY (report_id)
        );"""
    mod_report = """ALTER TABLE report AUTO_INCREMENT = 50000;"""
    mod_vuln_1 = """ALTER TABLE vulns ADD report_id int;"""
    mod_vuln_2 = """ALTER TABLE vulns ADD FOREIGN KEY (report_id) REFERENCES report(report_id);"""

    os.system('clear')
    banner()
    print """["""+warn+"""]Please make sure you have previously selected "(Re-)Initialize Server" in Gauntlet."""
    raw_input("["+question+"]Press enter to continue...")
    try:
        gauntlet = db.cursor()
        gauntlet.execute(create_table)
        gauntlet.execute(mod_report)
        gauntlet.execute(mod_vuln_1)
        gauntlet.execute(mod_vuln_2)
        gauntlet.close()
    except:
        print "\n["+warn+"]Please report this error to " + __maintainer__ + " by email at: " + __email__
        raise
    main_menu()

    print "["+note+"]You can now upload a new master dataset to Gauntlet"
    main_menu()


def pentest_checklist():
    """Generate a pentest checklist to be used for an assessment"""

    logging.info('Entered into pentest_checklist function')

    def build_html():
        """Generate HTML pentest checklist"""

        html_part = ''

        host_part = '\n<table>\n\t<tr id="$host-id" class="Host">\n\t\t<th class="Host-Header-Check"><input type=' \
                    '"checkbox"></th>\n\t\t<th class="Host-Header">$host-name</th>\n\t</tr>'
        port_part = '\n\t<tr id="$host-id_$port-id" class="Port">\n\t\t<td class="Port-Header-Check"><input ' \
                    'type="checkbox"></td>\n\t\t<td class="Port-Header">$port-text</td>\n\t</tr>'
        tool_part = '\n\t<tr id="$host-id_$port-id_$tool-id" class="Tool-$tool-name">\n\t\t<td colspan="2" class=' \
                    '"Tool-Header"title="$tool-name">$vuln-title</td>\n\t</tr>'
        tool_output_part = '\n\t<tr id="$host-id_$port-id_$tool-id_Output" class="Tool_Output">\n\t\t<td colspan="2">' \
                           '<pre>$tool-output</pre></td>\n\t</tr>'
        note_part = '\n\t<tr class="Notes">\n\t\t<td colspan="2"><input type="text" class="Notes-Text"></td>\n\t</tr>'

        for host in hosts2:
            # Build table Header for each host
            html_part += host_part
            if hosts2[host]['fqdn'][0] is not "":
                html_part = html_part.replace('$host-name', hosts2[host]['ipv4'] + " - " + hosts2[host]['fqdn'][0])
            else:
                html_part = html_part.replace('$host-name', hosts2[host]['ipv4'])

            # Build Port Rows
            for port_id in hosts2[host]['ports']:
                html_part += port_part

                # Build Vulnerability Rows
                if 'vulns' in hosts2[host]['ports'][port_id].keys() and args.tool_output:
                    for vuln_id in hosts2[host]['ports'][port_id]['vulns']:
                        html_part += tool_part
                        html_part += tool_output_part
                        tool_name = hosts2[host]['ports'][port_id]['vulns'][vuln_id]['tool']
                        html_part = html_part.replace('$tool-id', tool_name + "-" + vuln_id)
                        html_part = html_part.replace('$tool-name', tool_name)
                        html_part = html_part.replace('$vuln-title',
                                                      str(hosts2[host]['ports'][port_id]['vulns'][vuln_id]['title']))
                        html_part = html_part.replace('$tool-output',
                                                      (hosts2[host]['ports'][port_id]['vulns'][vuln_id]['output'])
                                                      .replace('&', '&amp;')
                                                      .replace('<', '&lt;')
                                                      .replace('>', '&gt;'))

                html_part = html_part.replace('$port-id', hosts2[host]['ports'][port_id]['port'] + "-" +
                                              hosts2[host]['ports'][port_id]['type'])
                html_part = html_part.replace('$port-text', hosts2[host]['ports'][port_id]['port'] + "/" +
                                              hosts2[host]['ports'][port_id]['type'] + "/" +
                                              str(hosts2[host]['ports'][port_id]['service']))

            html_part = html_part.replace('$host-id', hosts2[host]['ipv4'])
            html_part += (note_part * 3)
            html_part += '\n</table>'



        out_dir = get_path()
        checklist = os.path.join(out_dir, "Guinevere_"+assessment+"_checklist.html")
        html_file = open(checklist, 'w')
        # Build HTML File
        css = open(os.path.join(G_root, 'static', 'G-Checklist', 'G-Checklist.css'), 'r').read()
        html = open(os.path.join(G_root, 'static', 'G-Checklist', 'G-Checklist_Template.html'), 'r').read()
        html = html.replace('$ASSESSMENT', assessment)
        html = html.replace('$CSS', css)
        html = html.replace('$DATA', html_part)
        html_file.write(html)
        html_file.close()
        print "["+warn+"]Report saved to: " + checklist

    logging.info('Entered into build_html function within the pentest_checklist function')
    assessment = get_assessment("the assessment to create a pentest checklist for")
    banner()
    print "["+note+"]Building Pentest Checklist for " + assessment + "..."
    # hosts1 holds the record set returned from the SQL query
    hosts1 = db_query("""SELECT hosts.host_id, ip_address, fqdn, port_id, port, protocol, name """
                      """FROM hosts INNER JOIN ports ON hosts.host_id=ports.host_id WHERE port IS NOT NULL""",
                      assessment)
    host_ids = []   # A list to store retrieved host IDs
    hosts2 = {}     # A dictionary that holds the data used to print the checklist

    for row in hosts1:
        if row[0] not in host_ids:
            host_ids.append(row[0])

    for host_id in host_ids:
        hosts2[host_id] = {'ports': {}}    # Create a key in the hosts2 dictionary from the host id
        for row in hosts1:
            if host_id == row[0]:
                hosts2[host_id].update({'ipv4': row[1]})
                if 'fqdn' in hosts2[host_id].keys():
                    if row[2] not in hosts2[host_id]['fqdn']:
                        hosts2[host_id]['fqdn'].append(row[2])
                else:
                    hosts2[host_id].update({'fqdn': [row[2]]})
                if row[4] > 0:
                    hosts2[host_id]['ports'][row[3]] = {}
                    hosts2[host_id]['ports'][row[3]].update({'port': row[4]})
                    hosts2[host_id]['ports'][row[3]].update({'type': row[5]})
                    hosts2[host_id]['ports'][row[3]].update({'service': row[6]})

                #This check for the "tool_title" is here for backward compatability, remove when version are higher
                columns = db_query("""SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_SCHEMA =
                                    'gauntlet_%s' AND TABLE_NAME = 'vulnerabilities'""" % (assessment), assessment)
                if any('tool_title' in column for column in columns):
                    tool_data = db_query("""SELECT vuln_id, gnaat_id, tool, txt, tool_title FROM vulnerabilities
                                        WHERE host_id=%s and port_id=%s""" % (str(host_id), str(row[3])), assessment)
                    hosts2[host_id]['ports'][row[3]]['vulns'] = {}
                    for tool in tool_data:
                        hosts2[host_id]['ports'][row[3]]['vulns'].update({tool[0]: {'gnaat_id': tool[1],
                                                                                    'tool': tool[2],
                                                                                    'title': tool[4],
                                                                                    'output': tool[3],
                                                                                    'vuln_id': tool[0]}})
                else:
                    tool_data = db_query("""SELECT vuln_id, gnaat_id, tool, txt FROM vulnerabilities
                                        WHERE host_id= %s and port_id= %s""" % (str(host_id), str(row[3])), assessment)
                    if tool_data:
                        hosts2[host_id]['ports'][row[3]]['vulns'] = {}
                    for tool in tool_data:
                        if tool[1] != "":
                            title = db_query("SELECT title from vulns where gnaat_id=%s" % (tool[1]), 'GauntletData')
                        else:
                            title = None
                        if title is not None and len(title) > 0:
                            hosts2[host_id]['ports'][row[3]]['vulns'].update({tool[0]: {'gnaat_id': tool[1],
                                                                                        'tool': tool[2],
                                                                                        'title': title[0][0],
                                                                                        'output': tool[3],
                                                                                        'vuln_id': tool[0]}})
                        else:
                            hosts2[host_id]['ports'][row[3]]['vulns'].update({tool[0]: {'gnaat_id': tool[1],
                                                                                        'tool': tool[2],
                                                                                        'title': None,
                                                                                        'output': tool[3],
                                                                                        'vuln_id': tool[0]}})

    build_html()
    raw_input("["+note+"]Press enter to return to the main menu")

    main_menu()


def get_cvss(gnaat_id, version=2):
    """Get a CVSS object using the gnaat_id from the vulns table in the GauntletData database"""

    logging.info('Entered into get_cvss function')

    c = db_query("SELECT DISTINCT `cvss_access_vector`, `cvss_access_complexity`, `cvss_authentication`, "
                 "`cvss_confidentiality_impact`, `cvss_integrity_impact`, `cvss_availability_impact`, "
                 "`cvss_exploitability`, `cvss_remediation_level`, `cvss_report_confidence` FROM "
                 "vulns WHERE gnaat_id='%s';" % gnaat_id, "GauntletData")

    AV = None   # Access Vector
    AC = None   # Access Complexity
    Au = None   # Authentication
    C = None    # Confidentiality
    I = None    # Integrity
    A = None    # Availability
    E = None    # Exploitability
    RL = None   # Remediation Level
    RC = None   # Report Confidence

    # Access Vector
    if c[0][0] == 'remote':
        AV = 'A'
    elif c[0][0] == 'local':
        AV = 'L'
    elif c[0][0] == 'network':
        AV = 'N'

    # Access Complexity
    if c[0][1] == 'low':
        AC = 'L'
    elif c[0][1] == 'medium':
        AC = 'M'
    elif c[0][1] == 'high':
        AC = 'H'

    # Authentication
    if c[0][2] == 'not_required':
        Au = 'N'
    elif c[0][2] == 'required':
        Au = 'S'
    elif c[0][2] == 'multiple':
        Au = 'M'

    # Confidentiality
    if c[0][3] == 'none':
        C = 'N'
    elif c[0][3] == 'complete':
        C = 'C'
    elif c[0][3] == 'partial':
        C = 'P'

    # Integrity
    if c[0][4] == 'none':
        I = 'N'
    elif c[0][4] == 'complete':
        I = 'C'
    elif c[0][4] == 'partial':
        I = 'P'

    # Availability
    if c[0][5] == 'none':
        A = 'N'
    elif c[0][5] == 'complete':
        A = 'C'
    elif c[0][5] == 'partial':
        A = 'P'

    # Exploitability
    if c[0][6] == 'not_defined':
        E = 'ND'
    elif c[0][6] == 'unproven':
        E = 'U'
    elif c[0][6] == 'proof_of_concept':
        E = 'POC'
    elif c[0][6] == 'functional':
        E = 'F'
    elif c[0][6] == 'high':
        E = 'H'

    # Remediation Level
    if c[0][7] == 'not_defined':
        RL = 'ND'
    elif c[0][7] == 'official':
        RL = 'OF'
    elif c[0][7] == 'workaround':
        RL = 'W'
    elif c[0][7] == 'unavailable':
        RL = 'U'
    elif c[0][7] == 'temporary':
        RL = 'TF'

    # Report Confidence
    if c[0][8] == 'not_defined':
        RC = 'ND'
    elif c[0][8] == 'confirmed':
        RC = 'C'
    elif c[0][8] == 'unconfirmed':
        RC = 'UC'
    elif c[0][8] == 'uncorroborated':
        RC = 'UR'

    vector = 'AV:%s/AC:%s/Au:%s/C:%s/I:%s/A:%s/E:%s/RL:%s/RC:%s' % (AV, AC, Au, C, I, A, E, RL, RC)

    cvss = CVSS2(vector)

    return cvss


if args.debug:
    print "\n["+warn+"]Debug output enabled"
    logging.basicConfig(stream=sys.stdout, format='%(asctime)s\t%(levelname)s\t%(message)s',
                        datefmt='%Y-%m-%d %I:%M:%S %p', level=logging.DEBUG)  # Log to STDOUT
    raw_input("Press enter to continue...")


if __name__ == '__main__':
    try:
        main_menu()
    except KeyboardInterrupt:
        logging.info('User Interrupt! Quitting')
        print "\n["+warn+"]User Interrupt! Quitting...."
    except SystemExit:
        pass
    except:
        print "\n["+warn+"]Please report this error to " + __maintainer__ + " by email at: "+ __email__
        raise
